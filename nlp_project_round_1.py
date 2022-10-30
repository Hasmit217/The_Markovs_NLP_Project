#####***************When running first time comment out download statements*******************#######

#nltk.download('stopwords')
#nltk.download('wordnet')
#nltk.download('omw-1.4')
#nltk.download('averaged_perceptron_tagger')
#nltk.download('punkt')


#*****Removing figures, images, tables, headings from book******##
import docx
import re
doc = docx.Document()
def readtxt(filename):
    doc1 = docx.Document(filename)
    for para in doc1.paragraphs:
        newtext = para.text
        text = re.sub(r'[1-9] [A-Za-z ]+', '', newtext)
        text = re.sub(r'Figure [0-9].[0-9]:[A-Za-z .]+', '', text)
        text = re.sub(r'Table [0-9].[0-9]:[A-Za-z .]+', '', text)
        if len(text)>50:
            doc.add_paragraph(text)

readtxt('data.docx')
doc.save('preProcessedData.docx')


#**********************Tokenization********************##
import nltk
fullTokens = []
def tokenization(filename):
    doc1 = docx.Document(filename)
    for para in doc1.paragraphs:
        nltk_tokens = nltk.word_tokenize(para.text)
        for token in nltk_tokens:
            fullTokens.append(token)

tokenization('preProcessedData.docx')
with open("tokens.txt", "w",encoding='utf-8') as outfile:
    outfile.write("\n".join(fullTokens))


##*********************Lemmatization********************##
from nltk.stem import WordNetLemmatizer
wordnet_lemmatizer = WordNetLemmatizer()
punctuations="?:!.,;)(][}{*“”"
lemmaList = []
for word in fullTokens:
    if word not in punctuations:
        lemmaList.append(wordnet_lemmatizer.lemmatize(word))

with open("tokensAfterlemma.txt", "w",encoding='utf-8') as outfile:
    outfile.write("\n".join(lemmaList))


# ##*********************Stemming************************##
from nltk.stem import PorterStemmer
stemList = []
ps = PorterStemmer()
for w in lemmaList:
    stemList.append(ps.stem(w))

with open("tokensafterStem.txt", "w",encoding='utf-8') as outfile:
    outfile.write("\n".join(stemList))


# ##*************frequency distribution******************##
freq = {}
def freqCount(list):
    for token in list:
        if(token in freq):
            freq[token] +=1
        else:
            freq[token] = 1

freqCount(stemList)
file = open('freq.txt','w',encoding='utf-8')
for key, value in freq.items():
    file.write(key+" :"+str(value)+"\n")


#*******************bar plot for top 50 frequent words*******************#
import collections
import itertools

sorted_x = sorted(freq.items(), key=lambda kv: kv[1],reverse=True)
sorted_dict = collections.OrderedDict(sorted_x)

top = dict(itertools.islice(sorted_dict.items(), 50))
import matplotlib.pyplot as plt
sizes = list(top.values())
labels = list(top.keys())
plt.barh(labels,sizes)  
plt.yticks(fontsize=7.5)
plt.title("Word vs Frequency")
plt.savefig("BarPlot_Word_frequency_with_stop_word"+".png", bbox_inches='tight')
plt.show()


# ##*******************Word Cloud************************##
from wordcloud import WordCloud
wcloud = WordCloud().generate_from_frequencies(freq)
plt.figure()
plt.imshow(wcloud, interpolation="bilinear")
plt.axis("off")
plt.savefig("WordCloud_With_StopWords"+".png", bbox_inches='tight')
plt.show()


# # ##********************Remove Stop Words*******************##
from nltk.corpus import stopwords
stop_words = set(stopwords.words('english'))
filteredList = []
for w in stemList:
    if w not in stop_words:
        filteredList.append(w)
with open("tokensWithoutStopWords.txt", "w",encoding='utf-8') as outfile:
    outfile.write("\n".join(filteredList))



#*************************************************************#


# # ##*************Frequency Distribution After Removing Stopwords***************##
freq1 = {}
def freqCount(list):
    for token in list:
        if(token in freq1):
            freq1[token] +=1
        else:
            freq1[token] = 1
freqCount(filteredList)

#**************************************************************************#
import collections
import itertools

sorted_y = sorted(freq1.items(), key=lambda kv: kv[1],reverse=True)
sorted_dict1 = collections.OrderedDict(sorted_y)

top1 = dict(itertools.islice(sorted_dict1.items(), 50))
import matplotlib.pyplot as plt
sizes1 = list(top1.values())
labels1 = list(top1.keys())
plt.barh(labels1,sizes1)  
plt.yticks(fontsize=7.5)
plt.title("Word vs Frequency")
plt.savefig("BarPlot_Word_frequency_without_stop_word"+".png", bbox_inches='tight')
plt.show()


#*****************Word Cloud after removing stop words*************************#
wcloud = WordCloud().generate_from_frequencies(freq1)
plt.figure()
plt.imshow(wcloud, interpolation="bilinear")
plt.axis("off")
plt.savefig("WordCloud_without_stopWords"+".png", bbox_inches='tight')
plt.show()


# ##******************RelationShip b/w word Length and frequency*********************##
from collections import Counter
counts = Counter(len(word) for word in stemList)
wordLengths = []
frequencies = []
file = open('freqWithLength.txt','w',encoding='utf-8')
file.write(" Len     Freq")
for length in range(1, max(counts.keys()) + 1):
    file.write("\n")
    wordLengths.append(length)
    frequencies.append(counts.get(length,0))
    file.write(f'{length:4d} {counts.get(length, 0):6d}')

#********************Graph Plotting*****************************#

plt.barh(wordLengths,frequencies)  
plt.yticks(fontsize=7.5)
plt.title("Word length v/s Frequency")
plt.savefig("BarPlot_WordLength_frequency"+".png", bbox_inches='tight')
plt.show()


# ##********************POS Tagging************************************##
from nltk import tag
taggedTokens = tag.pos_tag(filteredList)
with open("taggedTokens.txt", "w",encoding='utf-8') as outfile:
    outfile.write("\n".join(" ".join(tup) for tup in taggedTokens))

##*********************Generating unique word file***********************************
uniq = {}
def freqCount(list):
    for token in list:
        if(token in uniq):
            uniq[token] +=1
        else:
            uniq[token] = 1
freqCount(stemList)
file = open('Unique.txt','w',encoding='utf-8')
for key, value in uniq.items():
    file.write(key+"\n")