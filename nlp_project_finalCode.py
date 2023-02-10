##### ***************When running first time comment out download statements*******************#######

# nltk.download('stopwords')
# nltk.download('wordnet')
# nltk.download('omw-1.4')
# nltk.download('averaged_perceptron_tagger')
# nltk.download('punkt')
import spacy
from sklearn.metrics import f1_score, accuracy_score
import numpy as np
# from spacy import tokenizer
# from spacy import displacy
import matplotlib.pyplot as plt


# # *****Removing figures, images, tables, headings from book******##
from nltk import tag
from nltk.corpus import stopwords, wordnet
from nltk.stem import PorterStemmer
from nltk.stem import WordNetLemmatizer
import nltk
import docx
import re
doc = docx.Document()


def readtxt(filename):
    doc1 = docx.Document(filename)
    for para in doc1.paragraphs:
        newtext = para.text
        text = re.sub(r'[1-9] [A-Za-z ]+', '', newtext)
        text = re.sub(r'Figure [0-9].[0-9]:[A-Za-z .]+', '', text)
        text = re.sub(r'Table [0-9].[0-9]:[A-Za-z .]+', '', text)
        if len(text) > 50:
            doc.add_paragraph(text)


readtxt('data.docx')
doc.save('preProcessedData.docx')


# # **********************Tokenization********************##
fullTokens = []


def tokenization(filename):
    doc1 = docx.Document(filename)
    for para in doc1.paragraphs:
        nltk_tokens = nltk.word_tokenize(para.text)
        for token in nltk_tokens:
            fullTokens.append(token)


tokenization('preProcessedData.docx')
with open("tokens.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(fullTokens))


# ## *********************Lemmatization********************##
wordnet_lemmatizer = WordNetLemmatizer()
punctuations = "?:!.,;)(][}{*“”"
lemmaList = []
for word in fullTokens:
    if word not in punctuations:
        lemmaList.append(wordnet_lemmatizer.lemmatize(word))

with open("tokensAfterlemma.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(lemmaList))


# # ##*********************Stemming************************##
stemList = []
ps = PorterStemmer()
for w in lemmaList:
    stemList.append(ps.stem(w))

with open("tokensafterStem.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(stemList))


# # ##*************frequency distribution******************##
freq = {}
def freqCount(list):
    for token in list:
        if(token in freq):
            freq[token] +=1
        else:
            freq[token] = 1

freqCount(stemList)
file = open('freq.txt','w',encoding='utf-8')
for key, value in freq.items():
    file.write(key+" :"+str(value)+"\n")


# # *******************bar plot for top 50 frequent words*******************#
import collections
import itertools

sorted_x = sorted(freq.items(), key=lambda kv: kv[1],reverse=True)
sorted_dict = collections.OrderedDict(sorted_x)

top = dict(itertools.islice(sorted_dict.items(), 50))
import matplotlib.pyplot as plt
sizes = list(top.values())
labels = list(top.keys())
plt.barh(labels,sizes)
plt.yticks(fontsize=7.5)
plt.title("Word vs Frequency")
plt.savefig("BarPlot_Word_frequency_with_stop_word"+".png", bbox_inches='tight')
plt.show()


# # ##*******************Word Cloud************************##
from wordcloud import WordCloud
wcloud = WordCloud().generate_from_frequencies(freq)
plt.figure()
plt.imshow(wcloud, interpolation="bilinear")
plt.axis("off")
plt.savefig("WordCloud_With_StopWords"+".png", bbox_inches='tight')
plt.show()


# # # ##********************Remove Stop Words*******************##
stop_words = set(stopwords.words('english'))
filteredList = []
for w in stemList:
    if w not in stop_words:
        filteredList.append(w)
with open("tokensWithoutStopWords.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(filteredList))


# # *************************************************************#


# # # ##*************Frequency Distribution After Removing Stopwords***************##
freq1 = {}
def freqCount(list):
    for token in list:
        if(token in freq1):
            freq1[token] +=1
        else:
            freq1[token] = 1
freqCount(filteredList)

# # **************************************************************************#
import collections
import itertools

sorted_y = sorted(freq1.items(), key=lambda kv: kv[1],reverse=True)
sorted_dict1 = collections.OrderedDict(sorted_y)

top1 = dict(itertools.islice(sorted_dict1.items(), 50))
import matplotlib.pyplot as plt
sizes1 = list(top1.values())
labels1 = list(top1.keys())
plt.barh(labels1,sizes1)
plt.yticks(fontsize=7.5)
plt.title("Word vs Frequency")
plt.savefig("BarPlot_Word_frequency_without_stop_word"+".png", bbox_inches='tight')
plt.show()


# # *****************Word Cloud after removing stop words*************************#
wcloud = WordCloud().generate_from_frequencies(freq1)
plt.figure()
plt.imshow(wcloud, interpolation="bilinear")
plt.axis("off")
plt.savefig("WordCloud_without_stopWords"+".png", bbox_inches='tight')
plt.show()


# # ##******************RelationShip b/w word Length and frequency*********************##
from collections import Counter
counts = Counter(len(word) for word in stemList)
wordLengths = []
frequencies = []
file = open('freqWithLength.txt','w',encoding='utf-8')
file.write(" Len     Freq")
for length in range(1, max(counts.keys()) + 1):
    file.write("\n")
    wordLengths.append(length)
    frequencies.append(counts.get(length,0))
    file.write(f'{length:4d} {counts.get(length, 0):6d}')

# # ********************Graph Plotting*****************************#

plt.barh(wordLengths,frequencies)
plt.yticks(fontsize=7.5)
plt.title("Word length v/s Frequency")
plt.savefig("BarPlot_WordLength_frequency"+".png", bbox_inches='tight')
plt.show()


# # ##********************POS Tagging************************************##
taggedTokens = tag.pos_tag(filteredList)
with open("taggedTokens.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(" ".join(tup) for tup in taggedTokens))

# # *********************Generating unique word file***********************************
uniq = {}


def freqCount(list):
    for token in list:
        if (token in uniq):
            uniq[token] += 1
        else:
            uniq[token] = 1


freqCount(stemList)
file = open('Unique.txt', 'w', encoding='utf-8')
for key, value in uniq.items():
    file.write(key+"\n")


# # *************************************POST_MID_PROJECT_CONTINUATION**********************************************

# # //////////////////////////Part-1/////////////////////////////////////////

taggedTokens1 = tag.pos_tag(fullTokens)
with open("taggedTokens1.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(" ".join(tup) for tup in taggedTokens1))

# ---------------------------------------------
nouns = []
for word, pos in taggedTokens1:
    if (pos == 'NN' or pos == 'NNP' or pos == 'NNS' or pos == 'NNPS'):
        if (nouns.count(word) == 0):
            nouns.append(word)

with open("nouns.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(nouns))
# ------------------------------------------------

verbs = []
for word, pos in taggedTokens1:
    if (pos == 'VB' or pos == 'VBD' or pos == 'VBG' or pos == 'VBN' or pos == 'VBZ'):
        if (verbs.count(word) == 0):
            verbs.append(word)

with open("verbs.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(verbs))
# --------------------------------------------------

nounlex = {}
for noun in nouns:
    syn = wordnet.synsets(noun)
    for s in syn:
        nounlex[noun] = s.lexname()
# -----------------------------------------------------
verblex = {}
for verb in verbs:
    syn = wordnet.synsets(verb)
    for s in syn:
        verblex[verb] = s.lexname()
# ---------------------------------------------------------

file = open('freqlex.txt', 'w', encoding='utf-8')
for key, value in nounlex.items():
    file.write(key+" :"+value+"\n")
for key, value in verblex.items():
    file.write(key+" :"+value+"\n")
# -----------------------------------------------------------

nounFreq = {}

for key, value in nounlex.items():
    if (value[0:4] == "noun"):
        if (value in nounFreq):
            nounFreq[value] += 1
        else:
            nounFreq[value] = 1
# -----------------------------------------------------------

verbFreq = {}

for key, value in verblex.items():
    if (value[0:4] == "verb"):
        if (value in verbFreq):
            verbFreq[value] += 1
        else:
            verbFreq[value] = 1

# -----------------------------------------------------------

plt.barh(list(nounFreq.keys()), nounFreq.values(), color='g')
plt.yticks(fontsize=7.5)
plt.title("Noun Categories v/s Frequency")
plt.savefig("BarPlot_NounCategories_Frequencies"+".png", bbox_inches='tight')
plt.show()

plt.barh(list(verbFreq.keys()), verbFreq.values(), color='g')
plt.yticks(fontsize=7.5)
plt.title("Verb Categories v/s Frequency")
plt.savefig("BarPlot_verbCategories_Frequencies"+".png", bbox_inches='tight')
plt.show()

# ///////////////////////Part-1 Complete////////////////////////////

# =======================================================================================

# //////////////////////Part-2//////////////////////////////////////

################ code for part2 NER #############

nlp = spacy.load('en_core_web_sm')

doc2 = docx.Document("preProcessedData.docx")
NER1 = []

for para in doc2.paragraphs:
    doc = nlp(para.text)
    for e in doc.ents:
        ents = e.text + " " + e.label_
        if (ents not in NER1):
            NER1.append(ents)
with open("NameEntity.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(NER1))

nounNER = []
for word in nouns:
    doc = nlp(word)
    for e in doc.ents:
        ents = e.text + " " + e.label_
        nounNER.append(ents)
# print(nounNER)
with open("nounNameEntity.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(nounNER))

verbNER = []
for word in verbs:
    doc = nlp(word)
    for e in doc.ents:
        ents = e.text + " " + e.label_
        verbNER.append(ents)
# print(verbNER)
with open("verbNameEntity.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(verbNER))
# ---------------------------------------------------------------------------------
fullTokens1 = []

def tokenization(filename):
    doc1 = docx.Document(filename)
    for para in doc1.paragraphs:
        nltk_tokens = nltk.word_tokenize(para.text)
        for token in nltk_tokens:
            fullTokens1.append(token)

tokenization('TestData.docx')
with open("tokensTest.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(fullTokens1))

taggedTokensTest1 = tag.pos_tag(fullTokens1)

doc2 = docx.Document("TestData.docx")

NER2 = []

for para in doc2.paragraphs:
    doc = nlp(para.text)
    for e in doc.ents:
        ents = e.text + " " + e.label_
        if (ents not in NER2):
            NER2.append(ents)

with open("NameEntityTest.txt", "w", encoding='utf-8') as outfile:
    outfile.write("\n".join(NER2))
# ---------------------------------------------------------------------------------

with open('predictedTest.txt', "r") as word_list:
    predictedList = word_list.read().split('\n')
print(predictedList)
with open('actualTest.txt', "r") as word_list:
    actualList = word_list.read().split('\n')
print(actualList)


# ----------------------------------Calculating F1 score and Accuracy----------------------------


f1 = f1_score(actualList, predictedList, average='micro')
print(f1)

acc = accuracy_score(actualList, predictedList, normalize=False)
print(acc)
